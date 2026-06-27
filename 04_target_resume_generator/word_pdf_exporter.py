"""
04_target_resume_generator/word_pdf_exporter.py — Word + PDF 双格式离线导出

依赖: python-docx (Word), reportlab (PDF)
功能: 一键导出 .docx 和 .pdf 简历文件，无需在线服务
"""

import json
from pathlib import Path
from datetime import datetime
from typing import Any


class WordExporter:
    """Word (.docx) 简历导出器"""

    @staticmethod
    def export(resume_data: dict, output_path: str | Path) -> str:
        """
        导出 Word 简历文档

        Args:
            resume_data: 简历数据（来自 TargetResumeGenerator.generate()）
            output_path: 输出文件路径

        Returns:
            输出文件的绝对路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(11)

        resume = resume_data.get("resume", resume_data)

        # ── 标题 ──
        title = doc.add_heading("个人简历", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── 基本信息 ──
        personal = resume.get("personal") or {}
        doc.add_heading("基本信息", level=1)
        info_items = [
            ("姓名", personal.get("name", "")),
            ("电话", personal.get("phone", "")),
            ("邮箱", personal.get("email", "")),
            ("城市", personal.get("city", "")),
            ("目标岗位", personal.get("target_job", "")),
            ("期望薪资", personal.get("salary", "")),
        ]
        for label, value in info_items:
            if value:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(value))

        # ── 个人概述 ──
        if resume.get("summary"):
            doc.add_heading("个人概述", level=1)
            doc.add_paragraph(resume["summary"])

        # ── 工作经历 ──
        experience = resume.get("experience") or []
        if experience:
            doc.add_heading("工作经历", level=1)
            for exp in experience:
                header = f"{exp.get('position', '')} | {exp.get('company', '')} | {exp.get('period', '')}"
                p = doc.add_paragraph()
                p.add_run(header).bold = True
                for bullet in exp.get("bullets") or []:
                    doc.add_paragraph(bullet, style="List Bullet")

        # ── 教育经历 ──
        education = resume.get("education") or []
        if education:
            doc.add_heading("教育经历", level=1)
            for edu in education:
                line = f"{edu.get('school', '')} | {edu.get('major', '')} | {edu.get('degree', '')} | {edu.get('period', '')}"
                doc.add_paragraph(line)

        # ── 技能 ──
        skills = resume.get("skills") or {}
        if skills:
            doc.add_heading("技能证书", level=1)
            for key, label in [("technical", "技术技能"), ("languages", "语言能力"), ("certificates", "证书")]:
                items = skills.get(key) or []
                if items:
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(", ".join(items))

        # ── 保存 ──
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        return str(path.absolute())


class PDFExporter:
    """PDF 简历导出器（使用 reportlab）"""

    @staticmethod
    def export(resume_data: dict, output_path: str | Path) -> str:
        """
        导出 PDF 简历文档

        Args:
            resume_data: 简历数据
            output_path: 输出文件路径

        Returns:
            输出文件的绝对路径
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            raise ImportError("请安装 reportlab: pip install reportlab")

        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=15*mm,
            bottomMargin=15*mm,
        )

        styles = getSampleStyleSheet()
        # 尝试注册中文字体（微软雅黑），失败则用默认字体
        try:
            pdfmetrics.registerFont(TTFont("SimSun", "simsun.ttc"))
            body_font = "SimSun"
        except Exception:
            body_font = "Helvetica"

        story = []
        resume = resume_data.get("resume", resume_data)

        # 标题
        title_style = ParagraphStyle("CTitle", parent=styles["Title"], alignment=TA_CENTER, fontName=body_font)
        story.append(Paragraph("个人简历", title_style))
        story.append(Spacer(1, 5*mm))

        # 基本信息
        personal = resume.get("personal") or {}
        story.append(Paragraph("<b>基本信息</b>", styles["Heading2"]))
        info_text = " | ".join(
            f"{k}: {v}"
            for k, v in [
                ("姓名", personal.get("name")),
                ("电话", personal.get("phone")),
                ("邮箱", personal.get("email")),
                ("城市", personal.get("city")),
                ("目标", personal.get("target_job")),
            ]
            if v
        )
        story.append(Paragraph(info_text, styles["Normal"]))
        story.append(Spacer(1, 3*mm))

        # 个人概述
        if resume.get("summary"):
            story.append(Paragraph("<b>个人概述</b>", styles["Heading2"]))
            story.append(Paragraph(resume["summary"], styles["Normal"]))
            story.append(Spacer(1, 3*mm))

        # 工作经历
        experience = resume.get("experience") or []
        if experience:
            story.append(Paragraph("<b>工作经历</b>", styles["Heading2"]))
            for exp in experience:
                header = f"{exp.get('position', '')} | {exp.get('company', '')} | {exp.get('period', '')}"
                story.append(Paragraph(f"<b>{header}</b>", styles["Normal"]))
                for bullet in exp.get("bullets") or []:
                    story.append(Paragraph(f"  • {bullet}", styles["Normal"]))
            story.append(Spacer(1, 3*mm))

        # 教育经历
        education = resume.get("education") or []
        if education:
            story.append(Paragraph("<b>教育经历</b>", styles["Heading2"]))
            for edu in education:
                line = f"{edu.get('school', '')} | {edu.get('major', '')} | {edu.get('degree', '')}"
                story.append(Paragraph(line, styles["Normal"]))

        # 技能
        skills = resume.get("skills") or {}
        if skills:
            story.append(Paragraph("<b>技能证书</b>", styles["Heading2"]))
            for key, label in [("technical", "技术"), ("languages", "语言"), ("certificates", "证书")]:
                items = skills.get(key) or []
                if items:
                    story.append(Paragraph(f"{label}: {', '.join(items)}", styles["Normal"]))

        doc.build(story)
        return str(path.absolute())


def export_resume(resume_data: dict, output_dir: str | Path, basename: str = "resume") -> dict:
    """
    一键导出 Word + PDF 双格式简历

    Args:
        resume_data: 简历数据
        output_dir: 输出目录
        basename: 文件名（不含扩展名）

    Returns:
        {"word": "path/to/file.docx", "pdf": "path/to/file.pdf"}
    """
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    word_path = WordExporter.export(resume_data, out / f"{basename}.docx")
    pdf_path = PDFExporter.export(resume_data, out / f"{basename}.pdf")

    return {"word": word_path, "pdf": pdf_path}


# ═══════════════════════════════════════════════════════════════
# 独立运行入口
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    # 测试导出
    test_data = {
        "resume": {
            "personal": {"name": "张三", "phone": "13800138000", "email": "test@qq.com", "city": "杭州", "target_job": "Python 开发"},
            "summary": "3 年 Python 后端开发经验，熟悉 Django/FastAPI。",
            "experience": [{"company": "A公司", "position": "Python 开发", "period": "2022-2024", "bullets": ["负责后端 API 开发"]}],
            "education": [{"school": "XX大学", "major": "计算机", "degree": "本科", "period": "2018-2022"}],
            "skills": {"technical": ["Python", "Django"], "languages": ["CET-6"], "certificates": []},
        }
    }
    result = export_resume(test_data, "./output_test")
    print(json.dumps(result, ensure_ascii=False, indent=2))
    print("[OK] Word + PDF 导出完成")
