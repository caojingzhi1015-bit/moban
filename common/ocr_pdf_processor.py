"""
common/ocr_pdf_processor.py — Unified file-to-text processor
============================================================
Handles TXT / PDF / DOCX / Image OCR for all business modules.

v2.0 — Robust cloud-ready extraction
  - PDF: pdfplumber per-page layout-aware extraction + PyPDF2 fallback
  - Scanned PDF detection → friendly prompt (no OCR on multi-page scans)
  - Image: PaddleOCR → pytesseract fallback
  - Section classification: regex + structural heuristics, bilingual
  - Section-formatted output for LLM prompts
  - Full exception isolation — one bad page never crashes the pipeline
  - No MinerU, no compiled C++ deps — stable on Streamlit Cloud
"""

from __future__ import annotations

import re
import io
import tempfile
from pathlib import Path
from typing import Callable
from dataclasses import dataclass, field


# ═══════════════════════════════════════════════════════════════
# Data types
# ═══════════════════════════════════════════════════════════════

@dataclass
class IndexedOutput:
    """Result of file processing — extracted text with metadata."""
    success: bool = True
    raw_text: str = ""              # plain extracted text
    indexed_text: str = ""          # line-numbered text: "[N] content"
    sectioned_text: str = ""        # section-organized with line indices (best for LLM)
    sections: dict = field(default_factory=dict)  # {key: {title, content: [...]}}
    method: str = "unknown"         # extraction method used
    page_count: int = 0
    is_scanned: bool = False        # True when PDF has no extractable text
    error: str = ""
    warnings: list = field(default_factory=list)


# ═══════════════════════════════════════════════════════════════
# Section detection patterns — comprehensive bilingual
# ═══════════════════════════════════════════════════════════════

# Each section has multiple regex patterns; a line matching any pattern
# (and shorter than 40 chars, preceded by blank line) is treated as section header.
SECTION_PATTERNS: dict[str, list[str]] = {
    "personal_info": [
        r"^(个人[信资][息料]|个人简介|个人概况|基[本础]信息|个人信息)$",
        r"^(Personal\s*Info(rmation)?|Contact\s*(Info|Details)?|Profile)$",
        r"^(姓名|Name|性别|Gender|出生|Birth|年龄|Age|籍贯|Native|民族|Nation|政治|Political)",
        r"^(电话|手机|Tel|Phone|Mobile|邮箱|Email|E-?mail|微信|WeChat|QQ|地址|Address)",
    ],
    "job_intent": [
        r"^(求职[意目][向标]|应聘[岗职][位称]|期望[岗职位][位称置]|目标[岗职位][位称置]|意向[岗职位])",
        r"^(期望薪资|期望[城地][市点]|目标薪资|到岗时间|可[到入]职)",
        r"^(Job\s*Objective|Career\s*Objective|Target\s*(Position|Job|Role)|Desired)",
    ],
    "education": [
        r"^(教育[经历][历背]?|教育程度|学历背景|教育$|学[习业]经历|学术背景|求学经历)",
        r"^(Education(al)?\s*(Background|Experience|History)?|Academic\s*(Background|History))$",
        r"^(毕业院校|学校名称|School|University|College|专业|Major|学位|Degree|GPA|绩点)",
    ],
    "work": [
        r"^(工作[经历][历背]?|工作履[历职]?|职业[经历][历生]?|从业[经历][历经]?|任职经历)$",
        r"^(Work\s*(Experience|History|Background)?|Employment|Professional\s*Experience|Career\s*History)$",
        r"^(工作单位|所在公司|雇主|Employer|Company)$",
    ],
    "project": [
        r"^(项目[经历][历背]?|项目经验|项目展示|项目列表|项目作品|核心项目|代表性项目|主要项目)$",
        r"^(Project(s)?\s*(Experience|History|Showcase|List)?|Portfolio|Key\s*Projects?)$",
    ],
    "internship": [
        r"^(实习[经历][历背]?|实习经验|Intern(ship)?\s*(Experience|History)?|Work\s*Placement)$",
    ],
    "skills": [
        r"^(技[能术](专长|特长|能[力方]面|栈|清单|列表)?|专业[技能][术力能]?|技术[能栈][力术]?|核心能力)$",
        r"^(Skills?|Technical\s+Skills?|Core\s+Competenc|Expertise|Technologies|Tech\s+Stack)$",
        r"^(掌握[技能术]|熟练掌握|精通|擅长|Proficien|Skilled\s*In)$",
    ],
    "certificates": [
        r"^(证书[及和列表]?|资格[证书][认证明]?|执业[资格证]|Certificates?|Certifications?|Licenses?|Credentials)$",
    ],
    "languages": [
        r"^(语言[能力][力水]?|外语[能力][力水]?|Languages?|Language\s+Skills?|Foreign\s+Languages?)$",
        r"^(英语[能力][力水]?|English\s*(Level|Proficiency|Skill))$",
    ],
    "self_intro": [
        r"^(自我[评价][价绍述]?|个人[评价][价绍述]?|个人优势|个人总结|自我总结|自我介绍|自我陈述)$",
        r"^(Self[- ]?(Assessment|Introduction|Description|Summary|Evaluation)|Personal\s+Statement|About\s+Me|Summary|Profile)$",
    ],
    "awards": [
        r"^(获[奖项][经历励]?|荣誉[奖项励]?|奖项|Awards?|Honors?|Achievements?|Accomplishments?|Scholarships?)$",
    ],
    "publications": [
        r"^(论文[发表]?|学术[成果][文果]?|专利|Publications?|Papers?|Patents?|Research\s*(Output|Work)?)$",
    ],
}


# ═══════════════════════════════════════════════════════════════
# Main processor class
# ═══════════════════════════════════════════════════════════════

class OcrPdfProcessor:
    """
    Unified file-to-text processor.
    Single entry point: process_file() — auto-detects type, extracts, classifies.

    Output:
      - raw_text: continuous plain text
      - indexed_text: each line prefixed with "[N]" for AI source tracing
      - sectioned_text: sections demarcated with 【Section Name】 headers + line indices
      - sections: structured dict for programmatic consumption
    """

    # ── Public entry ──────────────────────────────────────────

    @staticmethod
    def process_file(
        file_path: str | Path,
        lang: str = "zh",
        on_progress: Callable | None = None,
    ) -> IndexedOutput:
        """
        Main entry point. Auto-detects file type and extracts text.

        Args:
            file_path: path to the file (PDF / image / docx / txt / ...)
            lang: "zh" or "en" — affects OCR language and section patterns
            on_progress: optional callback(status_message: str)

        Returns IndexedOutput — check .success before consuming .raw_text etc.
        """
        path = Path(file_path)
        if not path.exists():
            return IndexedOutput(success=False, error=f"文件不存在: {file_path}")

        suffix = path.suffix.lower()
        result_text = ""
        method = "unknown"
        warnings: list[str] = []
        page_count = 0
        is_scanned = False

        try:
            # ── Plain text ──
            if suffix in (".txt", ".md", ".py", ".json", ".csv",
                          ".log", ".xml", ".html", ".yaml", ".yml"):
                result_text = _read_text_file(path)
                method = "text"

            # ── PDF ──
            elif suffix == ".pdf":
                outcome = _read_pdf(path, lang, on_progress)
                result_text = outcome["text"]
                method = outcome["method"]
                page_count = outcome.get("page_count", 0)
                is_scanned = outcome.get("is_scanned", False)
                warnings = outcome.get("warnings", [])

            # ── DOCX / DOC ──
            elif suffix in (".docx", ".doc"):
                result_text = _read_docx(path)
                method = "docx"

            # ── Images ──
            elif suffix in (".png", ".jpg", ".jpeg", ".webp",
                            ".bmp", ".gif", ".tiff", ".tif"):
                result_text = _ocr_image(path, lang)
                method = "ocr"
                if not result_text.strip():
                    result_text = _ocr_image_tesseract(path, lang)
                    if result_text.strip():
                        method = "ocr_tesseract"

            # ── Unknown — try as text ──
            else:
                result_text = _read_text_file(path)
                method = "text_fallback"

        except MemoryError:
            return IndexedOutput(
                success=False,
                error="文件过大导致内存不足，请尝试压缩后重新上传",
            )
        except PermissionError:
            return IndexedOutput(
                success=False,
                error="文件读取权限不足，请检查文件是否被其他程序占用",
            )
        except Exception as e:
            return IndexedOutput(
                success=False,
                error=f"文件解析异常 [{type(e).__name__}]: {e}",
            )

        # ── Post-extraction checks ──
        if not result_text or not result_text.strip():
            if is_scanned:
                return IndexedOutput(
                    success=False,
                    is_scanned=True,
                    page_count=page_count,
                    method=method,
                    error=(
                        "⚠️ 当前PDF为扫描件/图片型PDF，无法自动提取文字。\n"
                        "请将简历正文直接粘贴到左侧文本框，或使用手机拍照后上传PNG/JPG图片。"
                    ),
                    warnings=warnings,
                )
            return IndexedOutput(
                success=False,
                error=(
                    "未能从文件中提取到任何文字内容。"
                    "文件可能为空、已损坏、或为不支持格式的图片。"
                ),
                warnings=warnings,
            )

        # ── Normalize ──
        result_text = _normalize_text(result_text)

        # ── Build indexed text (every line numbered for AI traceability) ──
        lines = result_text.strip().split("\n")
        indexed = "\n".join(f"[{i}] {line}" for i, line in enumerate(lines))

        # ── Classify sections ──
        sections = _classify_sections(result_text, lang)

        # ── Build sectioned text (【Section】 + indexed content) ──
        sectioned_text = _format_sectioned_text(sections, result_text, lang)

        return IndexedOutput(
            success=True,
            raw_text=result_text.strip(),
            indexed_text=indexed,
            sectioned_text=sectioned_text,
            sections=sections,
            method=method,
            page_count=page_count,
            is_scanned=is_scanned,
            warnings=warnings,
        )

    # ── Public utility ────────────────────────────────────────

    @staticmethod
    def text_to_indexed_lines(text: str) -> str:
        """Convert plain text to line-numbered format for AI prompt."""
        return "\n".join(
            f"[{i}] {line}" for i, line in enumerate(text.split("\n"))
        )

    # ── Backward-compatible wrappers (delegate to module functions) ──

    @staticmethod
    def _classify_sections(text: str, lang: str = "zh") -> dict:
        """Public wrapper for section classification."""
        return _classify_sections(text, lang)

    @staticmethod
    def _format_sectioned_text(sections: dict, raw_text: str, lang: str = "zh") -> str:
        """Public wrapper for sectioned text formatting."""
        return _format_sectioned_text(sections, raw_text, lang)


# ═══════════════════════════════════════════════════════════════
# File-type readers (module-level for picklability)
# ═══════════════════════════════════════════════════════════════

def _read_text_file(path: Path) -> str:
    """Read a plain-text file with encoding fallback."""
    for enc in ("utf-8", "gbk", "gb2312", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=enc, errors="replace")
        except Exception:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


# ── PDF ──────────────────────────────────────────────────────

def _read_pdf(path: Path, lang: str = "zh",
              on_progress: Callable | None = None) -> dict:
    """
    Multi-level PDF text extraction.
    Level 1: pdfplumber — per-page, layout-aware
    Level 2: PyPDF2 — plain text fallback
    Level 3: detect scanned → return empty with is_scanned=True

    Returns: {text, method, page_count, is_scanned, warnings}
    """
    warnings: list[str] = []
    page_count = 0

    # ── Level 1: pdfplumber (layout-aware, per-page isolation) ──
    try:
        import pdfplumber
        all_parts: list[str] = []
        empty_pages = 0

        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = _extract_page_text(page)
                    if page_text and page_text.strip():
                        all_parts.append(page_text)
                    else:
                        empty_pages += 1
                except Exception as e:
                    warnings.append(f"第{page_num}页提取异常: {e}")
                    # Last-resort: basic extract_text
                    try:
                        t = page.extract_text()
                        if t and t.strip():
                            all_parts.append(t)
                        else:
                            empty_pages += 1
                    except Exception:
                        empty_pages += 1

            if on_progress:
                on_progress(f"pdfplumber: {len(all_parts)}/{page_count} 页已提取")

        text = "\n\n".join(all_parts)

        if text.strip() and _is_meaningful_text(text):
            if empty_pages > page_count * 0.5:
                warnings.append(
                    f"{empty_pages}/{page_count} 页几乎无文字（可能包含扫描页）"
                )
            return {
                "text": text,
                "method": "pdfplumber",
                "page_count": page_count,
                "is_scanned": False,
                "warnings": warnings,
            }
        elif text.strip():
            warnings.append("pdfplumber 提取的文字质量较低")
    except ImportError:
        warnings.append("pdfplumber 未安装，尝试 PyPDF2")
    except Exception as e:
        warnings.append(f"pdfplumber 打开失败: {e}")

    # ── Level 2: PyPDF2 ──
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        if page_count == 0:
            page_count = len(reader.pages)
        parts = []
        for page in reader.pages:
            try:
                t = page.extract_text()
                if t:
                    parts.append(t)
            except Exception:
                pass
        text = "\n".join(parts)
        if text.strip() and _is_meaningful_text(text):
            return {
                "text": text,
                "method": "pypdf2",
                "page_count": page_count,
                "is_scanned": False,
                "warnings": warnings,
            }
    except ImportError:
        warnings.append("PyPDF2 未安装")
    except Exception as e:
        warnings.append(f"PyPDF2 失败: {e}")

    # ── Level 3: scanned PDF — no text extractable ──
    return {
        "text": "",
        "method": "none",
        "page_count": page_count,
        "is_scanned": True,
        "warnings": warnings,
    }


def _extract_page_text(page) -> str:
    """
    Extract text from a single pdfplumber page with layout awareness.

    Strategy:
      1. Try extract_text(layout=True) — preserves horizontal spacing
      2. If text appears multi-column (many short lines), use word-level
         column detection to reorder
      3. Fall back to plain extract_text()
    """
    # Primary: layout-preserving extraction
    text = None
    try:
        text = page.extract_text(layout=True)
    except Exception:
        pass

    if not text or not text.strip():
        try:
            text = page.extract_text()
        except Exception:
            return ""

    if not text or not text.strip():
        return ""

    # Multi-column detection heuristic:
    # If > 30% of lines are < 30 chars AND the page width is wide (> 400pt),
    # the page likely has columns that got interleaved.
    lines = text.split("\n")
    short_lines = sum(1 for l in lines if len(l.strip()) < 30 and l.strip())
    short_ratio = short_lines / max(1, len(lines))

    page_width = getattr(page, "width", 0)

    if short_ratio > 0.3 and page_width > 400:
        # Try column-aware extraction via word clustering
        try:
            col_text = _extract_page_columns(page)
            if col_text and len(col_text) > len(text) * 0.5:
                return col_text
        except Exception:
            pass

    return text


def _extract_page_columns(page) -> str:
    """
    Extract text using word-level positions to detect and reorder columns.

    Clusters words by x-coordinate into column groups,
    then outputs each column top-to-bottom, columns left-to-right.
    """
    try:
        words = page.extract_words(
            x_tolerance=3,
            y_tolerance=3,
            keep_blank_chars=False,
        )
    except Exception:
        return page.extract_text() or ""

    if not words or len(words) < 3:
        return page.extract_text() or ""

    # Cluster into columns by x-coordinate gaps
    columns = _cluster_columns(words)

    if len(columns) <= 1:
        return page.extract_text() or ""

    # Render each column
    col_texts = []
    for col_words in columns:
        col_words.sort(key=lambda w: (w["top"], w["x0"]))
        col_texts.append(_words_to_lines(col_words))

    return "\n".join(col_texts)


def _cluster_columns(words: list[dict]) -> list[list[dict]]:
    """Cluster words into columns based on horizontal gaps."""
    if len(words) < 2:
        return [words]

    # All x-center positions, sorted
    x_centers = sorted((w["x0"] + w["x1"]) / 2 for w in words)

    # Find significant gaps
    gaps = []
    for i in range(1, len(x_centers)):
        gap = x_centers[i] - x_centers[i - 1]
        if gap > 0:
            gaps.append((gap, x_centers[i - 1], x_centers[i]))

    if not gaps:
        return [words]

    # Median gap as baseline; column breaks are 3× median gap (min 30pt)
    median_gap = sorted(g[0] for g in gaps)[len(gaps) // 2]
    threshold = max(median_gap * 3, 30)

    # Find column dividers
    split_xs = [round((x1 + x2) / 2, 1) for gap, x1, x2 in gaps if gap > threshold]

    if not split_xs:
        return [words]

    split_xs.sort()
    columns = [[] for _ in range(len(split_xs) + 1)]

    for w in words:
        x_center = (w["x0"] + w["x1"]) / 2
        col_idx = sum(1 for sx in split_xs if x_center > sx)
        columns[col_idx].append(w)

    return [c for c in columns if c]


def _words_to_lines(words: list[dict]) -> str:
    """Group sorted words into text lines by vertical proximity."""
    if not words:
        return ""

    lines = []
    current_line_words = [words[0]["text"]]
    current_top = words[0]["top"]
    tolerance = 5  # px — words within 5px vertically are on the same line

    for w in words[1:]:
        if abs(w["top"] - current_top) > tolerance:
            lines.append(" ".join(current_line_words))
            current_line_words = [w["text"]]
            current_top = w["top"]
        else:
            current_line_words.append(w["text"])

    if current_line_words:
        lines.append(" ".join(current_line_words))

    return "\n".join(lines)


# ── Image OCR ────────────────────────────────────────────────

def _ocr_image(path: Path, lang: str = "zh") -> str:
    """OCR an image using PaddleOCR."""
    try:
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(lang="ch" if lang == "zh" else "en")
        result = ocr.ocr(str(path))
        if result and result[0]:
            lines = [line[1][0] for line in result[0] if line]
            return "\n".join(lines)
    except ImportError:
        pass
    except Exception:
        pass
    return ""


def _ocr_image_tesseract(path: Path, lang: str = "zh") -> str:
    """OCR an image using pytesseract as fallback."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(path)
        text = pytesseract.image_to_string(
            img,
            lang="chi_sim+eng" if lang == "zh" else "eng",
        )
        return text
    except ImportError:
        pass
    except Exception:
        pass
    return ""


# ── DOCX ─────────────────────────────────────────────────────

def _read_docx(path: Path) -> str:
    """Extract text from Word documents with table support."""
    try:
        from docx import Document
        doc = Document(str(path))
        parts = []

        # Iterate document body in order (paragraphs + tables)
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                texts = []
                for t_node in element.findall(".//w:t", nsmap):
                    if t_node.text:
                        texts.append(t_node.text)
                line = "".join(texts).strip()
                if line:
                    parts.append(line)

            elif tag == "tbl":
                # Table — extract row by row
                for row in element.findall(".//w:tr", nsmap):
                    cell_texts = []
                    for cell in row.findall(".//w:tc", nsmap):
                        cell_parts = []
                        for p in cell.findall(".//w:p", nsmap):
                            p_texts = [
                                t.text or ""
                                for t in p.findall(".//w:t", nsmap)
                            ]
                            cell_parts.append("".join(p_texts))
                        cell_texts.append(" ".join(cell_parts).strip())
                    row_text = " | ".join(c for c in cell_texts if c)
                    if row_text.strip():
                        parts.append(row_text)

        text = "\n".join(parts)
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback: simple paragraph-only extraction
    try:
        from docx import Document
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""


# ═══════════════════════════════════════════════════════════════
# Text normalization
# ═══════════════════════════════════════════════════════════════

def _normalize_text(text: str) -> str:
    """Clean extracted text: fix line endings, collapse blanks, remove artifacts."""
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove form-feed characters
    text = text.replace("\f", "\n")

    # Collapse multiple spaces (preserve newlines)
    text = re.sub(r"[^\S\n]{2,}", " ", text)

    # Merge 4+ consecutive blank lines → 3
    text = re.sub(r"\n{4,}", "\n\n\n", text)

    # Fix common PDF hyphenation artifacts (word broken across lines)
    text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)

    # Remove common header/footer patterns (page numbers, repetitive headers)
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*第\s*\d+\s*页\s*(共\s*\d+\s*页)?\s*$", "", text,
                  flags=re.MULTILINE)
    text = re.sub(r"^\s*Page\s+\d+\s*(of\s+\d+)?\s*$", "", text,
                  flags=re.MULTILINE | re.IGNORECASE)

    # Remove completely empty lines that remain
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text


def _is_meaningful_text(text: str) -> bool:
    """Check if extracted text has real content (not garbled/sparse)."""
    stripped = text.strip()

    if len(stripped) < 20:
        return False

    # Printable character ratio
    printable = sum(1 for c in stripped if c.isprintable() or c in "\n\t ")
    if printable / max(1, len(stripped)) < 0.8:
        return False

    # At least some letters/numbers/CJK characters
    meaningful = sum(
        1 for c in stripped
        if c.isalnum() or ("一" <= c <= "鿿") or ("぀" <= c <= "ヿ")
    )
    if meaningful / max(1, len(stripped)) < 0.25:
        return False

    return True


# ═══════════════════════════════════════════════════════════════
# Section classification
# ═══════════════════════════════════════════════════════════════

def _classify_sections(text: str, lang: str = "zh") -> dict:
    """
    Split text into hierarchical sections using regex + structural heuristics.

    Returns:
      {section_key: {"title": str, "content": [line, ...]}, ...}

    Keys: header, personal_info, job_intent, education, work, project,
          internship, skills, certificates, languages, self_intro, awards,
          publications
    """
    lines = text.split("\n")
    sections: dict[str, dict] = {}
    current_section = "header"
    section_content: list[str] = []

    # Compile all patterns
    compiled: dict[str, list[re.Pattern]] = {}
    for key, patterns in SECTION_PATTERNS.items():
        compiled[key] = [re.compile(p, re.IGNORECASE) for p in patterns]

    def _match_section(line: str) -> str | None:
        """Return section key if line matches a section-header pattern."""
        stripped = line.strip()
        if not stripped:
            return None
        # Section headers are short
        if len(stripped) > 40:
            return None
        for key, pats in compiled.items():
            for pat in pats:
                if pat.search(stripped):
                    return key
        return None

    def _finalize(sec_name: str, content: list[str], title: str = ""):
        """Save accumulated content lines into the named section."""
        nonlocal sections
        if sec_name not in sections:
            sections[sec_name] = {"title": title, "content": []}
        clean = [l for l in content if l.strip()]
        if clean:
            if sections[sec_name]["content"]:
                sections[sec_name]["content"].extend(clean)
            else:
                sections[sec_name]["content"] = clean
            if title and not sections[sec_name]["title"]:
                sections[sec_name]["title"] = title

    prev_blank = True  # treat start of doc as after blank line

    for line in lines:
        stripped = line.strip()

        # Blank line — potential section break
        if not stripped:
            prev_blank = True
            section_content.append("")
            continue

        # Try to match as section header
        detected = _match_section(stripped)

        if detected and prev_blank and len(stripped) >= 2:
            # Section header detected — save previous section, start new
            _finalize(current_section, section_content)
            current_section = detected
            section_content = [stripped]  # keep header line in content
            prev_blank = False
            continue

        # Also detect numbered section headers: "一、", "1.", "（一）", "①"
        if prev_blank and len(stripped) < 35 and re.match(
            r"^[（(]?[一二三四五六七八九十\d]{1,3}[）).、\s]",
            stripped,
        ):
            # Could be a subsection — keep in current section as-is
            pass

        # Regular content line
        section_content.append(stripped)
        prev_blank = False

    # Finalize last section
    _finalize(current_section, section_content)

    # Post-process: if "header" has very few lines, merge into next section
    if "header" in sections:
        header_content = sections["header"].get("content", [])
        # Clean blank lines
        header_content = [l for l in header_content if l.strip()]
        if len(header_content) <= 4:
            # Find first non-header section and prepend header content
            for key in sections:
                if key != "header" and sections[key].get("content"):
                    sections[key]["content"] = header_content + sections[key]["content"]
                    break
            # Remove header to avoid duplication in output
            del sections["header"]

    return sections


def _format_sectioned_text(sections: dict, raw_text: str, lang: str = "zh") -> str:
    """
    Build LLM-friendly sectioned text with line indices preserved.

    Format:
        【基本信息 / Personal Information】
        [3] 张三
        [4] 电话: 13800138000

        【教育经历 / Education】
        [8] 2018-2022  北京大学  计算机科学

    This helps the LLM understand document structure and
    reference source line numbers for traceability.
    """
    if not sections:
        # No sections detected — return simple indexed text
        lines = raw_text.strip().split("\n")
        return "\n".join(f"[{i}] {line}" for i, line in enumerate(lines))

    # Build line index lookup: for each line of raw_text, what's its index?
    raw_lines = raw_text.strip().split("\n")

    # Section display labels
    SECTION_LABELS: dict[str, tuple[str, str]] = {
        "personal_info": ("基本信息", "Personal Information"),
        "job_intent": ("求职意向", "Job Intention"),
        "education": ("教育经历", "Education"),
        "work": ("工作经历", "Work Experience"),
        "project": ("项目经历", "Projects"),
        "internship": ("实习经历", "Internship"),
        "skills": ("专业技能", "Skills"),
        "certificates": ("证书资格", "Certificates"),
        "languages": ("语言能力", "Languages"),
        "self_intro": ("自我评价", "Self-Introduction"),
        "awards": ("获奖荣誉", "Awards & Honors"),
        "publications": ("论文专利", "Publications & Patents"),
        "header": ("文档头部", "Document Header"),
    }

    parts: list[str] = []

    for sec_key in sections:
        sec = sections[sec_key]
        content = sec.get("content", [])
        # Filter blanks
        content = [l for l in content if l.strip()]
        if not content:
            continue

        # Find line indices for each content line
        indexed_content: list[str] = []
        for cl in content:
            # Try to find this line's index in raw_lines
            idx = _find_line_index(raw_lines, cl)
            if idx >= 0:
                indexed_content.append(f"[{idx}] {cl}")
            else:
                indexed_content.append(f"[?] {cl}")

        zh_label, en_label = SECTION_LABELS.get(
            sec_key, (sec_key, sec_key)
        )
        label = f"{zh_label} / {en_label}" if lang == "zh" else f"{en_label}"

        parts.append(f"\n{'─' * 50}")
        parts.append(f"【{label}】")
        parts.append(f"{'─' * 50}")
        parts.append("\n".join(indexed_content))

    return "\n".join(parts) if parts else OcrPdfProcessor.text_to_indexed_lines(raw_text)


def _find_line_index(raw_lines: list[str], target: str) -> int:
    """Find the index of a target string in raw_lines. Returns -1 if not found."""
    target_stripped = target.strip()
    if not target_stripped:
        return -1

    # Exact match
    for i, line in enumerate(raw_lines):
        if line.strip() == target_stripped:
            return i

    # Substring match (target is part of a longer line)
    for i, line in enumerate(raw_lines):
        if target_stripped in line:
            return i

    return -1
