import sys, json, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

src = r"D:\xwechat_files\wxid_0fozf2qgws8o22_c057\msg\file\2026-06\程序使用说明.docx"
dst = r"D:\KuGou\Lyric\简介生成 - 副本\作业提交\程序使用说明.docx"

# 直接从源模板创建新文档
doc = Document(src)
# 清除所有段落
for p in doc.paragraphs:
    p.clear()

def H(text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold
    r.font.size = Pt(size); r.font.name = "Microsoft YaHei"

def P(text, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = "Microsoft YaHei"

def C(text, size=10, gray=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = "Microsoft YaHei"
    if gray: r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def S(code):
    for line in code.split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(9); r.font.name = "Consolas"

def B(text, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(size); r.font.name = "Microsoft YaHei"

# ===== 文档内容 =====
H("程序使用说明", 16)
H("CareerAI \xb7 求职助手", 12, bold=False)
doc.add_paragraph()

P("本程序为解决求职者在求职过程中面临的简历优化、职位匹配、面试准备等痛点而开发。程序包含6个功能模块，覆盖从JD解析到模拟面试的完整求职链路。")
doc.add_paragraph()

B("功能模块图", 14)
doc.add_paragraph()

B("模块1 \u2014\u2014 JD职位解析器", 12)
C("本模块中涉及到的第三方模块：re（正则表达式）、json（JSON处理）", 10, gray=True)
doc.add_paragraph()
B("2.1 模块功能", 12)
P("本模块可以实现招聘JD文本的结构化解析，提取职位名称、公司名、工作地点、薪资范围、学历要求、年限要求、技能关键词等信息。")
doc.add_paragraph()
B("JD解析功能", 11)
P("输入JD文本可提取结构化职位信息，主要使用正则表达式模式匹配技术（30+技能类别），配合AI融合提取策略。核心代码块：")
S("""    @classmethod
    def extract_jd(cls, text: str) -> dict:
        skills = []
        for pat in cls.HARD_SKILL_PATTERNS:
            for m in pat.finditer(text):
                if m.group(0) not in skills:
                    skills.append(m.group(0))
        return {"position":...,"hard_skills":skills,...}

    merge_results(): AI优先、正则兜底、技能去重""")
doc.add_paragraph()

B("模块2 \u2014\u2014 个人简历解析器", 12)
C("本模块中涉及到的第三方模块：re（正则表达式）、json（JSON处理）", 10, gray=True)
doc.add_paragraph()
B("2.2 模块功能", 12)
P("本模块可以实现简历文本的解析，提取姓名、手机号、邮箱、城市、教育经历、工作经历、技能等硬字段。主要通过正则表达式提取强格式字段，再用merge_results()合并AI结果。")
doc.add_paragraph()
B("简历解析功能", 11)
P("核心代码：")
S("""    @staticmethod
    def extract_hard_fields(text: str) -> dict:
        # 姓名/电话/邮箱/城市/学校/专业/公司/时间/技能 8类硬字段
        name_m = re.search(r"\\u59d3\\s*\\u540d[\\uff1a:\\s]*([^\\n,...")
        phone_m = re.search(r"(1[3-9]\\d{9})", text)
        email_m = re.search(r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\\.[a-z]...)")
        school_pat = re.compile(r"((?:[\\u4e00-\\u9fff]{2,8})(?:大学|学院))")
        return {"name":name,"phone":phone,"email":email,"schools":[],...}

    to_material_json(): 标准化素材库导出格式""")
doc.add_paragraph()

B("模块3 \u2014\u2014 AI职位匹配追问系统", 12)
C("本模块中涉及到的第三方模块：re（正则表达式）、json（JSON处理）", 10, gray=True)
B("2.3 模块功能", 12)
P("本模块实现JD与简历之间的6维度缺口分析（技能/软技能/量化/项目/年限/意向），并生成针对性的追问问题。使用规则模板引擎确保问题精准有效。")
doc.add_paragraph()
B("缺口分析与追问生成", 11)
S("""    @staticmethod
    def analyze_gaps(jd_data, resume_data):
        # 6维度分析: skill_gap/soft_skill/quant/project/exp/job_target
        for skill in jd_data.get("hard_skills",[]):
            if skill.lower() not in all_text:
                gaps.append({"category":"skill_gap","keyword":skill})
        return gaps

    build_template_questions(): 规则模板引擎生成追问""")
doc.add_paragraph()

B("模块4 \u2014\u2014 JD对标简历生成器", 12)
C("本模块中涉及到的第三方模块：re（正则表达式）、json（JSON处理）", 10, gray=True)
B("2.4 模块功能", 12)
P("本模块根据JD要求对简历进行重组优化，生成五模块结构化简历（个人信息/个人概述/工作经历/教育经历/技能证书），同时生成中英文自我介绍文本并估算朗读时长。")
doc.add_paragraph()
B("简历生成功能", 11)
S("""    核心技术：
    _by_jd_relevance(): JD关键词权重排序算法
    五模块构建: _build_personal/_build_summary/_build_experience/...
    build_self_intro(): 中英文自我介绍生成+朗读时长估算

    @staticmethod
    def _by_jd_relevance(entries, jd):
        keywords=jd.get("hard_skills",[])+jd.get("soft_skills",[])
        return sorted(entries,key=lambda e:sum(1 for k in keywords
            if k.lower() in json.dumps(e).lower()), reverse=True)""")
doc.add_paragraph()

B("模块5 \u2014\u2014 AI模拟面试官系统", 12)
C("本模块中涉及到的第三方模块：json（JSON处理）", 10, gray=True)
B("2.5 模块功能", 12)
P("本模块实现4阶段结构化面试（开场\\u2192背景深挖\\u2192能力匹配\\u2192收尾），支持4种面试官人格（HR/技术/压力/英文）。使用4阶段状态机管理面试流程，本地规则引擎基于简历数据模板化生成追问（无需API），可自动生成面试评估报告。")
doc.add_paragraph()
B("面试对话功能", 11)
S("""    4阶段状态机: opening\\u2192background\\u2192competency\\u2192closing
    _build_local_reply(): 纯规则引擎（无需API）基于简历生成追问
    4种人格: hr/tech/stress/english
    generate_report(): 面试评估报告自动生成""")
doc.add_paragraph()

B("模块6 \u2014\u2014 多模型API统一调度网关", 12)
C("本模块中涉及到的第三方模块：time（时间处理）、json、collections", 10, gray=True)
B("2.6 模块功能", 12)
P("本模块实现多模型API统一调度网关，包含令牌桶限流器（RateLimiter）控制请求频率、用量追踪器（UsageTracker）统计Token消耗、余额守护器（BillingGuard）预警预算超限、按任务类型分流模型选择（lite/pro）。其他5个模块全部调用本网关发起请求。")
doc.add_paragraph()
B("限流与用量统计", 11)
S("""    RateLimiter: 令牌桶限流算法（控制RPM）
    UsageTracker: API用量追踪（token数/请求数/成本）
    BillingGuard: 余额阈值守护（预警+封顶）
    select_model(): 按任务类型分流（lite vs pro）""")
doc.add_paragraph()

# 学习感受
B("本学期学习感受", 14)
doc.add_paragraph()
P("通过《计算思维之问题求解》课程的学习，我深刻理解了如何将计算思维应用于实际问题解决。本课程让我掌握了Python程序设计的基本方法，学会了如何将复杂问题分解为可执行的模块化解决方案。在开发CareerAI求职助手的过程中，我综合运用了正则表达式、数据结构、算法设计、模块化编程等技术，体会到了计算思维在真实场景中的强大应用价值。课程中关于问题分解、模式识别、抽象化和算法设计的思想，不仅帮助我完成了这个项目，也让我在解决其他实际问题时有了更清晰的思路。")

doc.save(dst)
print("OK: done")
